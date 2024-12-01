import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
import plotly.io as pio
from typing import Dict, List, Optional, Union, Any
import base64
import json
import yaml
import os
from pathlib import Path
import logging
from datetime import datetime
import tempfile
import csv
import sqlite3
from io import BytesIO, StringIO
import zipfile
import xlsxwriter

logger = logging.getLogger(__name__)

class DataExporter:
    """Class to handle data export to various formats"""
    
    def __init__(self, export_dir: str = "exports/"):
        """Initialize Data Exporter"""
        self.export_dir = export_dir
        self._init_exporter()
        
    def _init_exporter(self):
        """Initialize export system"""
        try:
            os.makedirs(self.export_dir, exist_ok=True)
            logger.info("Export system initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing export system: {str(e)}")
            raise

    def export_to_excel(self,
                       data: Union[pd.DataFrame, Dict[str, pd.DataFrame]],
                       filename: str,
                       **kwargs) -> BytesIO:
        """
        Export data to Excel file
        
        Args:
            data: DataFrame or dict of DataFrames for multiple sheets
            filename: Output filename
            **kwargs: Additional export options
        """
        try:
            output = BytesIO()
            
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                # Handle single DataFrame
                if isinstance(data, pd.DataFrame):
                    data = {'Sheet1': data}
                
                # Write each sheet
                for sheet_name, df in data.items():
                    df.to_excel(
                        writer,
                        sheet_name=sheet_name,
                        index=kwargs.get('include_index', True)
                    )
                    
                    # Auto-adjust columns width
                    worksheet = writer.sheets[sheet_name]
                    for idx, col in enumerate(df.columns):
                        max_length = max(
                            df[col].astype(str).str.len().max(),
                            len(str(col))
                        )
                        worksheet.set_column(idx, idx, max_length + 2)
                
                # Add metadata if provided
                if 'metadata' in kwargs:
                    metadata_sheet = writer.book.add_worksheet('Metadata')
                    for idx, (key, value) in enumerate(kwargs['metadata'].items()):
                        metadata_sheet.write(idx, 0, key)
                        metadata_sheet.write(idx, 1, str(value))
            
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error exporting to Excel: {str(e)}")
            raise

    def export_to_csv(self,
                     data: pd.DataFrame,
                     filename: str,
                     **kwargs) -> StringIO:
        """Export data to CSV file"""
        try:
            output = StringIO()
            
            data.to_csv(
                output,
                index=kwargs.get('include_index', True),
                encoding=kwargs.get('encoding', 'utf-8'),
                sep=kwargs.get('separator', ','),
                quoting=kwargs.get('quoting', csv.QUOTE_MINIMAL)
            )
            
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error exporting to CSV: {str(e)}")
            raise

    def export_to_json(self,
                      data: Union[pd.DataFrame, Dict],
                      filename: str,
                      **kwargs) -> str:
        """Export data to JSON file"""
        try:
            if isinstance(data, pd.DataFrame):
                result = data.to_json(
                    orient=kwargs.get('orient', 'records'),
                    date_format=kwargs.get('date_format', 'iso')
                )
            else:
                result = json.dumps(
                    data,
                    indent=kwargs.get('indent', 2),
                    default=str
                )
            
            return result
            
        except Exception as e:
            logger.error(f"Error exporting to JSON: {str(e)}")
            raise

    def export_to_sql(self,
                     data: pd.DataFrame,
                     table_name: str,
                     db_path: str,
                     **kwargs) -> bool:
        """Export data to SQLite database"""
        try:
            conn = sqlite3.connect(db_path)
            
            data.to_sql(
                table_name,
                conn,
                if_exists=kwargs.get('if_exists', 'replace'),
                index=kwargs.get('include_index', True)
            )
            
            conn.close()
            return True
            
        except Exception as e:
            logger.error(f"Error exporting to SQL: {str(e)}")
            raise

    def export_to_parquet(self,
                         data: pd.DataFrame,
                         filename: str,
                         **kwargs) -> BytesIO:
        """Export data to Parquet file"""
        try:
            output = BytesIO()
            
            data.to_parquet(
                output,
                engine=kwargs.get('engine', 'auto'),
                compression=kwargs.get('compression', 'snappy')
            )
            
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error exporting to Parquet: {str(e)}")
            raise

    def export_to_html(self,
                      data: pd.DataFrame,
                      filename: str,
                      **kwargs) -> str:
        """Export data to HTML file"""
        try:
            return data.to_html(
                index=kwargs.get('include_index', True),
                classes=kwargs.get('classes', 'table table-striped'),
                escape=kwargs.get('escape', True)
            )
            
        except Exception as e:
            logger.error(f"Error exporting to HTML: {str(e)}")
            raise

    def export_to_word(self,
                      content: Dict[str, Any],
                      filename: str,
                      **kwargs) -> BytesIO:
        """Export content to Word document"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(content.get('title', 'Report'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            if 'metadata' in content:
                doc.add_heading('Metadata', level=1)
                for key, value in content['metadata'].items():
                    doc.add_paragraph(f"{key}: {value}")
            
            # Add content sections
            for section in content.get('sections', []):
                doc.add_heading(section['title'], level=1)
                doc.add_paragraph(section['content'])
                
                # Add table if present
                if 'table' in section:
                    table_df = pd.DataFrame(section['table'])
                    table = doc.add_table(
                        rows=len(table_df)+1,
                        cols=len(table_df.columns)
                    )
                    
                    # Add header
                    for j, column in enumerate(table_df.columns):
                        table.cell(0, j).text = str(column)
                    
                    # Add data
                    for i in range(len(table_df)):
                        for j, value in enumerate(table_df.iloc[i]):
                            table.cell(i+1, j).text = str(value)
                
                # Add figure if present
                if 'figure' in section:
                    doc.add_picture(
                        section['figure'],
                        width=Inches(6)
                    )
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            raise

    def export_to_pdf(self,
                     content: Dict[str, Any],
                     filename: str,
                     **kwargs) -> bytes:
        """Export content to PDF file"""
        try:
            # Create HTML content
            html_content = f"""
            <html>
                <head>
                    <style>
                        body {{ font-family: Arial, sans-serif; }}
                        .container {{ max-width: 800px; margin: 0 auto; padding: 20px; }}
                        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                        th, td {{ padding: 8px; border: 1px solid #ddd; }}
                        th {{ background-color: #f5f5f5; }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <h1>{content.get('title', 'Report')}</h1>
            """
            
            # Add metadata
            if 'metadata' in content:
                html_content += "<h2>Metadata</h2>"
                html_content += "<table>"
                for key, value in content['metadata'].items():
                    html_content += f"<tr><th>{key}</th><td>{value}</td></tr>"
                html_content += "</table>"
            
            # Add content sections
            for section in content.get('sections', []):
                html_content += f"<h2>{section['title']}</h2>"
                html_content += f"<p>{section['content']}</p>"
                
                if 'table' in section:
                    df = pd.DataFrame(section['table'])
                    html_content += df.to_html()
                
                if 'figure' in section:
                    html_content += f'<img src="data:image/png;base64,{section["figure"]}">'
            
            html_content += """
                    </div>
                </body>
            </html>
            """
            
            # Convert HTML to PDF
            options = {
                'page-size': 'A4',
                'margin-top': '20mm',
                'margin-right': '20mm',
                'margin-bottom': '20mm',
                'margin-left': '20mm',
            }
            
            return pdfkit.from_string(html_content, False, options=options)
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            raise

    def create_download_link(self,
                           data: Union[bytes, str, BytesIO, StringIO],
                           filename: str,
                           mime_type: str) -> str:
        """Create download link for exported data"""
        try:
            if isinstance(data, (BytesIO, StringIO)):
                data = data.getvalue()
            if isinstance(data, str):
                data = data.encode()
                
            b64 = base64.b64encode(data).decode()
            return f'<a href="data:{mime_type};base64,{b64}" download="{filename}">Download {filename}</a>'
            
        except Exception as e:
            logger.error(f"Error creating download link: {str(e)}")
            raise

    def show_export_interface(self):
        """Show export interface in Streamlit"""
        st.subheader("📥 Data Export")
        
        if not st.session_state.uploaded_files:
            st.warning("Please upload some data files first!")
            return
        
        # Select dataset
        selected_file = st.selectbox(
            "Select dataset to export:",
            list(st.session_state.uploaded_files.keys())
        )
        
        if selected_file:
            df = st.session_state.uploaded_files[selected_file]['data']
            
            # Export options
            export_format = st.selectbox(
                "Export Format",
                ["Excel", "CSV", "JSON", "SQL", "Parquet", "HTML", "Word", "PDF"]
            )
            
            # Common options
            include_index = st.checkbox("Include Index", value=True)
            
            if export_format == "Excel":
                if st.button("Export to Excel"):
                    output = self.export_to_excel(
                        df,
                        f"{selected_file}.xlsx",
                        include_index=include_index
                    )
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.xlsx",
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        ),
                        unsafe_allow_html=True
                    )
                    
            elif export_format == "CSV":
                separator = st.selectbox("Separator", [",", ";", "\t", "|"])
                encoding = st.selectbox("Encoding", ["utf-8", "utf-16", "ascii"])
                
                if st.button("Export to CSV"):
                    output = self.export_to_csv(
                        df,
                        f"{selected_file}.csv",
                        include_index=include_index,
                        separator=separator,
                        encoding=encoding
                    )
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.csv",
                            "text/csv"
                        ),
                        unsafe_allow_html=True
                    )
                    
            elif export_format == "JSON":
                orient = st.selectbox(
                    "JSON Orient",
                    ["records", "columns", "index", "split", "table"]
                )
                indent = st.number_input("Indent", min_value=0, value=2)
                
                if st.button("Export to JSON"):
                    output = self.export_to_json(
                        df,
                        f"{selected_file}.json",
                        orient=orient,
                        indent=indent
                    )
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.json",
                            "application/json"
                        ),
                        unsafe_allow_html=True
                    )
                    
            elif export_format == "SQL":
                table_name = st.text_input("Table Name", value=selected_file.split('.')[0])
                if_exists = st.selectbox("If Table Exists", ["replace", "append", "fail"])
                
                if st.button("Export to SQL"):
                    db_path = os.path.join(self.export_dir, f"{selected_file}.db")
                    if self.export_to_sql(
                        df,
                        table_name,
                        db_path,
                        if_exists=if_exists,
                        include_index=include_index
                    ):
                        st.success(f"Data exported to SQLite database: {db_path}")
                        
            elif export_format == "Parquet":
                compression = st.selectbox("Compression", ["snappy", "gzip", "brotli"])
                
                if st.button("Export to Parquet"):
                    output = self.export_to_parquet(
                        df,
                        f"{selected_file}.parquet",
                        compression=compression
                    )
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.parquet",
                            "application/octet-stream"
                        ),
                        unsafe_allow_html=True
                    )

            elif export_format == "HTML":
                classes = st.text_input(
                    "CSS Classes",
                    value="table table-striped"
                )
                escape = st.checkbox("Escape HTML", value=True)
                
                if st.button("Export to HTML"):
                    output = self.export_to_html(
                        df,
                        f"{selected_file}.html",
                        include_index=include_index,
                        classes=classes,
                        escape=escape
                    )
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.html",
                            "text/html"
                        ),
                        unsafe_allow_html=True
                    )
                    
            elif export_format == "Word":
                # Prepare Word document content
                st.write("### Document Content")
                
                title = st.text_input("Document Title", value="Data Analysis Report")
                
                # Add metadata
                with st.expander("Add Metadata"):
                    metadata = {}
                    num_metadata = st.number_input("Number of Metadata Fields", min_value=0, value=3)
                    for i in range(num_metadata):
                        col1, col2 = st.columns(2)
                        with col1:
                            key = st.text_input(f"Key {i+1}")
                        with col2:
                            value = st.text_input(f"Value {i+1}")
                        if key and value:
                            metadata[key] = value
                
                # Add sections
                sections = []
                num_sections = st.number_input("Number of Sections", min_value=1, value=2)
                
                for i in range(num_sections):
                    with st.expander(f"Section {i+1}"):
                        section = {
                            'title': st.text_input(f"Section {i+1} Title"),
                            'content': st.text_area(f"Section {i+1} Content")
                        }
                        
                        # Add table
                        if st.checkbox(f"Include Table in Section {i+1}"):
                            section['table'] = df
                        
                        # Add visualization if available
                        if 'visualizations' in st.session_state:
                            viz_options = [f"Visualization {j+1}" for j in range(len(st.session_state.visualizations))]
                            selected_viz = st.selectbox(f"Add Visualization to Section {i+1}", ["None"] + viz_options)
                            
                            if selected_viz != "None":
                                viz_idx = int(selected_viz.split()[-1]) - 1
                                img_bytes = pio.to_image(st.session_state.visualizations[viz_idx]['figure'], format='png')
                                section['figure'] = BytesIO(img_bytes)
                        
                        sections.append(section)
                
                if st.button("Export to Word"):
                    content = {
                        'title': title,
                        'metadata': metadata,
                        'sections': sections
                    }
                    
                    output = self.export_to_word(
                        content,
                        f"{selected_file}.docx"
                    )
                    
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        ),
                        unsafe_allow_html=True
                    )
                    
            elif export_format == "PDF":
                # Similar options as Word export
                st.write("### Document Content")
                
                title = st.text_input("Document Title", value="Data Analysis Report")
                
                # Add metadata
                with st.expander("Add Metadata"):
                    metadata = {}
                    num_metadata = st.number_input("Number of Metadata Fields", min_value=0, value=3)
                    for i in range(num_metadata):
                        col1, col2 = st.columns(2)
                        with col1:
                            key = st.text_input(f"Key {i+1}")
                        with col2:
                            value = st.text_input(f"Value {i+1}")
                        if key and value:
                            metadata[key] = value
                
                # Add sections
                sections = []
                num_sections = st.number_input("Number of Sections", min_value=1, value=2)
                
                for i in range(num_sections):
                    with st.expander(f"Section {i+1}"):
                        section = {
                            'title': st.text_input(f"Section {i+1} Title"),
                            'content': st.text_area(f"Section {i+1} Content")
                        }
                        
                        # Add table
                        if st.checkbox(f"Include Table in Section {i+1}"):
                            section['table'] = df
                        
                        # Add visualization if available
                        if 'visualizations' in st.session_state:
                            viz_options = [f"Visualization {j+1}" for j in range(len(st.session_state.visualizations))]
                            selected_viz = st.selectbox(f"Add Visualization to Section {i+1}", ["None"] + viz_options)
                            
                            if selected_viz != "None":
                                viz_idx = int(selected_viz.split()[-1]) - 1
                                img_bytes = pio.to_image(st.session_state.visualizations[viz_idx]['figure'], format='png')
                                section['figure'] = base64.b64encode(img_bytes).decode()
                        
                        sections.append(section)
                
                if st.button("Export to PDF"):
                    content = {
                        'title': title,
                        'metadata': metadata,
                        'sections': sections
                    }
                    
                    output = self.export_to_pdf(
                        content,
                        f"{selected_file}.pdf"
                    )
                    
                    st.markdown(
                        self.create_download_link(
                            output,
                            f"{selected_file}.pdf",
                            "application/pdf"
                        ),
                        unsafe_allow_html=True
                    )

if __name__ == "__main__":
    exporter = DataExporter()
    exporter.show_export_interface()